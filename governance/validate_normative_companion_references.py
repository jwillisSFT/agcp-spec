#!/usr/bin/env python3
from __future__ import annotations

import hashlib
import json
import re
import sys
import zipfile
from datetime import date
from pathlib import Path

from docx import Document
from openpyxl import load_workbook

ROOT = Path(__file__).resolve().parents[1]
DISPOSITION_PATH = ROOT / "governance/normative-companion-reference-dispositions.json"
REPORT_PATH = ROOT / "governance/AGCP-normative-companion-reference-validation.json"

TEXT_SUFFIXES = {".md", ".json", ".yaml", ".yml", ".csv", ".txt", ".tex"}
OFFICE_SUFFIXES = {".docx", ".xlsx", ".pptx"}
ALLOWED_RETIRED_LABEL_PATHS = {
    Path("governance/normative-companion-reference-dispositions.json"),
    Path("governance/AGCP-normative-companion-reference-validation.json"),
}

EXPECTED_BINDINGS = {
    "spec/AGCP-Human-Review-Specification.md": [
        "DS-020 Governance Evidence",
        "../schemas/governance_evidence.json",
        "DS-033 Evidence Qualification Result",
        "../schemas/evidence_qualification_result.json",
    ],
    "spec/ledger/AGCP-Append-Only-Governance-Ledger-Specification.md": [
        "../../schemas/governance_evidence.json",
        "../../schemas/evidence_qualification_result.json",
        "../AGCP-Multitenant-Operational-Specification.md",
        "../AGCP-Error-Mapping.md",
    ],
    "conformance/AGCP-Conformance.md": [
        "AGCP Multitenant Operational Specification",
        "AGCP Human Adjudication and Governance Approval Specification",
        "DS-020 Governance Evidence",
        "DS-033 Evidence Qualification Result",
    ],
    "lifecycle/AGCP Normative Governance Progression Table.md": [
        "../schemas/governance_evidence.json",
        "../schemas/evidence_qualification_result.json",
        "../spec/AGCP-Human-Review-Specification.md",
    ],
    "lifecycle/AGCP Governance Progression Implementation Guide.md": [
        "../schemas/governance_evidence.json",
        "../schemas/evidence_qualification_result.json",
        "../spec/AGCP-Provenance-Wire-Format-Specification.md",
        "../spec/AGCP-Multitenant-Operational-Specification.md",
        "../spec/AGCP-Error-Mapping.md",
    ],
    "governance/AGCP-Versioning.md": [
        "Normative reference integrity",
        "AGCP-Normative-Companion-Reference-Dispositions.md",
        "../schemas/governance_evidence.json",
        "../schemas/evidence_qualification_result.json",
    ],
    "lifecycle/AGCP Governance Lifecycle Model.md": [
        "../schemas/governance_evidence.json",
        "../schemas/evidence_qualification_result.json",
        "../spec/AGCP-Error-Mapping.md",
        "../spec/AGCP-Human-Review-Specification.md",
    ],
}


def sha256(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as f:
        for chunk in iter(lambda: f.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def text_from_office(path: Path) -> str:
    if path.suffix.lower() == ".docx":
        doc = Document(path)
        parts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    parts.extend(p.text for p in cell.paragraphs)
        return "\n".join(parts)
    if path.suffix.lower() == ".xlsx":
        wb = load_workbook(path, read_only=True, data_only=False)
        parts: list[str] = []
        for ws in wb.worksheets:
            for row in ws.iter_rows():
                for cell in row:
                    if cell.value is not None:
                        parts.append(str(cell.value))
        return "\n".join(parts)
    if path.suffix.lower() == ".pptx":
        # Office XML fallback is sufficient for repository-reference strings.
        with zipfile.ZipFile(path) as zf:
            xml = "\n".join(
                zf.read(name).decode("utf-8", "ignore")
                for name in zf.namelist()
                if name.endswith(".xml")
            )
        return re.sub(r"<[^>]+>", "", xml)
    return ""


def add_check(checks: list[dict], check_id: str, passed: bool, detail: str) -> None:
    checks.append({"check_id": check_id, "status": "PASS" if passed else "FAIL", "detail": detail})


def main() -> int:
    checks: list[dict] = []
    issues: list[str] = []

    disposition = json.loads(DISPOSITION_PATH.read_text(encoding="utf-8"))
    dispositions = disposition.get("dispositions", [])
    add_check(checks, "NCR-001", len(dispositions) == 3, f"Controlled dispositions: {len(dispositions)}")
    if len(dispositions) != 3:
        issues.append("Expected three controlled reference dispositions.")

    retired_labels: list[str] = []
    replacement_paths: set[str] = set()
    for item in dispositions:
        retired_labels.extend(item.get("retired_labels", []))
        replacement_paths.update(item.get("replacement_artifacts", []))

    missing_replacements = sorted(p for p in replacement_paths if not (ROOT / p).exists())
    add_check(
        checks,
        "NCR-002",
        not missing_replacements,
        "All replacement artifacts exist." if not missing_replacements else f"Missing: {missing_replacements}",
    )
    issues.extend(f"Missing replacement artifact: {p}" for p in missing_replacements)

    active_hits: list[dict] = []
    scanned_text = 0
    scanned_office = 0
    scope_counts = {"specifications": 0, "openapi": 0, "registries": 0, "examples": 0, "catalogs": 0, "rtm_office": 0}

    for path in sorted(ROOT.rglob("*")):
        if not path.is_file() or ".git" in path.parts:
            continue
        rel = path.relative_to(ROOT)
        if rel in ALLOWED_RETIRED_LABEL_PATHS:
            continue
        content = ""
        if path.suffix.lower() in TEXT_SUFFIXES:
            try:
                content = path.read_text(encoding="utf-8")
            except UnicodeDecodeError:
                continue
            scanned_text += 1
        elif path.suffix.lower() in OFFICE_SUFFIXES:
            content = text_from_office(path)
            scanned_office += 1
        else:
            continue
        for label in retired_labels:
            if label in content:
                active_hits.append({"path": str(rel), "label": label})
        rels = str(rel)
        if rels.startswith(("spec/", "lifecycle/", "conformance/", "governance/")):
            scope_counts["specifications"] += 1
        if rels == "api/AGCP-HTTP-Contract.yaml" or rels.startswith("api/"):
            scope_counts["openapi"] += 1
        if rels.startswith("registries/"):
            scope_counts["registries"] += 1
        if rels.startswith("schemas/examples/"):
            scope_counts["examples"] += 1
        if "catalog" in rel.name.lower():
            scope_counts["catalogs"] += 1
        if rels.endswith("AGCP_Requirements_Traceability_Matrix_(RTM).xlsx"):
            scope_counts["rtm_office"] += 1

    add_check(
        checks,
        "NCR-003",
        not active_hits,
        f"Scanned {scanned_text} text artifacts and {scanned_office} Office artifacts; active retired-label hits: {len(active_hits)}",
    )
    for hit in active_hits:
        issues.append(f"Active retired reference {hit['label']!r} in {hit['path']}")

    missing_bindings: list[str] = []
    for rel, tokens in EXPECTED_BINDINGS.items():
        text = (ROOT / rel).read_text(encoding="utf-8")
        for token in tokens:
            if token not in text:
                missing_bindings.append(f"{rel}: {token}")
    add_check(
        checks,
        "NCR-004",
        not missing_bindings,
        "All seven corrected documents contain their canonical replacement bindings."
        if not missing_bindings
        else f"Missing bindings: {missing_bindings}",
    )
    issues.extend(f"Missing canonical binding: {x}" for x in missing_bindings)

    canonical_title = "AGCP Human Adjudication and Governance Approval Specification"
    human_review = (ROOT / "spec/AGCP-Human-Review-Specification.md").read_text(encoding="utf-8")
    add_check(checks, "NCR-005", human_review.startswith(f"# {canonical_title}"), "Canonical human-review title verified.")
    if not human_review.startswith(f"# {canonical_title}"):
        issues.append("Canonical human-review title is absent.")

    # Verify referenced local paths embedded in the corrected files.
    unresolved_paths: list[str] = []
    code_path_pattern = re.compile(r"`([^`]+\.(?:md|json|yaml|yml|csv|docx|xlsx))`")
    for rel in EXPECTED_BINDINGS:
        source = ROOT / rel
        text = source.read_text(encoding="utf-8")
        for target in code_path_pattern.findall(text):
            if target.startswith(("http://", "https://")):
                continue
            # Only explicit relative links are resolved here. Repository-root
            # path literals are verified through the controlled disposition
            # replacement-artifact existence check.
            if not target.startswith(("../", "./")):
                continue
            resolved = (source.parent / target).resolve()
            if not resolved.exists():
                unresolved_paths.append(f"{rel} -> {target}")
    add_check(
        checks,
        "NCR-006",
        not unresolved_paths,
        "All canonical local path references in corrected documents resolve."
        if not unresolved_paths
        else f"Unresolved: {unresolved_paths}",
    )
    issues.extend(f"Unresolved canonical path: {x}" for x in unresolved_paths)

    # Required scope confirmation explicitly includes the sources named by P1-01.
    required_scope = all(scope_counts[k] > 0 for k in scope_counts)
    add_check(checks, "NCR-007", required_scope, f"Repository scope scanned: {scope_counts}")
    if not required_scope:
        issues.append(f"One or more required repository scopes were not scanned: {scope_counts}")

    source_files = sorted(set(EXPECTED_BINDINGS) | {
        "governance/AGCP-Normative-Companion-Reference-Dispositions.md",
        "governance/normative-companion-reference-dispositions.json",
        "spec/README.md",
        "lifecycle/README.md",
        "README.md",
        "ARCHITECTURE.md",
        "conformance/agcp-conformance-manifest.yml",
    })
    source_hashes = {rel: sha256(ROOT / rel) for rel in source_files if (ROOT / rel).exists()}

    report = {'release_context':{'repository_release_target':'v2.0.4','repository_release_target_status':'PUBLIC_REVIEW_CONTROLLED_BASELINE','controlling_published_baseline':'v2.0.4','controlling_baseline_status':'PUBLIC_REVIEW_CONTROLLED_BASELINE','baseline_date':'2026-08-05','artifact_lifecycle_state':'CURRENT'},
        "validation_id": "AGCP-NORMATIVE-COMPANION-REFERENCE-VALIDATION",
        "finding": "P1-01",
        "release_target": "v2.0.4",
        "validation_date": str(date.today()),
        "status": "PASS" if not issues else "FAIL",
        "summary": {
            "checks": len(checks),
            "passed": sum(1 for c in checks if c["status"] == "PASS"),
            "failed": sum(1 for c in checks if c["status"] == "FAIL"),
            "text_artifacts_scanned": scanned_text,
            "office_artifacts_scanned": scanned_office,
            "retired_labels": len(retired_labels),
            "active_retired_reference_hits": len(active_hits),
            "replacement_artifacts": len(replacement_paths),
            "corrected_documents": len(EXPECTED_BINDINGS),
        },
        "scope_counts": scope_counts,
        "checks": checks,
        "issues": issues,
        "source_hashes": source_hashes,
    }
    REPORT_PATH.write_text(json.dumps(report, indent=2) + "\n", encoding="utf-8")
    print(json.dumps({"status": report["status"], "checks": len(checks), "issues": issues}, indent=2))
    return 0 if not issues else 1


if __name__ == "__main__":
    sys.exit(main())
