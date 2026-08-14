#!/usr/bin/env python3
"""Validate the AGCP v2 controlled Normative Statement inventory disposition."""
from __future__ import annotations
import hashlib, json, re, sys
from pathlib import Path
from docx import Document
from openpyxl import load_workbook
from release_version import ROOT, RELEASE_TAG, RTM_SPEC_VERSION, CURRENT_RELEASE_NOTES, release_context

REPORT=ROOT/"governance/AGCP-normative-statement-inventory-validation.json"
NS_DOC=ROOT/"spec/AGCP Normative Statements.docx"
CORE=ROOT/"spec/AGCP-Core.docx"
RTM=ROOT/"spec/AGCP_Requirements_Traceability_Matrix_(RTM).xlsx"
DUPLICATE_TEXT="Conformance to this specification requires implementation of all mandatory normative requirements applicable to the claimed implementation profile."
EXPECTED_UNIQUE=357; EXPECTED_OCCURRENCES=358; CANONICAL="NS-2.7-01"; ALIAS="NS-17.2-01"

def sha(p): return hashlib.sha256(p.read_bytes()).hexdigest()
def ids_from_doc(doc):
    pat=re.compile(r"^NS-\d+(?:\.\d+)*[A-Z]?-\d+$")
    return {p.text.strip() for p in doc.paragraphs if pat.fullmatch(p.text.strip())}
def split_ids(v): return [x.strip() for x in str(v or '').split(';') if x.strip() and x.strip().upper()!='N/A']

def main():
    checks=[]; issues=[]
    def ck(name,ok,detail):
        checks.append({"check":name,"status":"PASS" if ok else "FAIL","detail":detail})
        if not ok: issues.append(f"{name}: {detail}")
    core=Document(CORE); core_hits=[i for i,p in enumerate(core.paragraphs) if p.text.strip()==DUPLICATE_TEXT]
    ck("duplicated_core_obligation_occurs_twice",len(core_hits)==2,{"paragraph_indexes":core_hits,"count":len(core_hits)})
    nsdoc=Document(NS_DOC); nsids=ids_from_doc(nsdoc)
    ck("normative_statement_document_unique_identifier_count",len(nsids)==EXPECTED_UNIQUE,{"unique_identifiers":len(nsids)})
    ck("canonical_identifier_present_and_alias_unassigned",CANONICAL in nsids and ALIAS not in nsids,{"canonical_present":CANONICAL in nsids,"alias_present":ALIAS in nsids})
    note="NS-17.2-01 is intentionally unassigned"
    ck("normative_statement_document_contains_disposition_note",any(note in p.text for p in nsdoc.paragraphs),note)
    wb=load_workbook(RTM,data_only=False); rel=wb["NS_CR_Relationships"]; rtmids=set()
    for row in rel.iter_rows(values_only=True):
        for value in row:
            if isinstance(value,str): rtmids.update(re.findall(r"NS-\d+(?:\.\d+)*[A-Z]?-\d+",value))
    # The relationship sheet can contain contextual examples; compare controlled unique set intersection through all RTM cells.
    primary=wb[wb.sheetnames[0]]; hdr={primary.cell(1,c).value:c for c in range(1,primary.max_column+1)}; primary_ids=set()
    for r in range(2,primary.max_row+1):
        for col in ["NS_ID","Normative_Statement_IDs","Test_Generating_NS_IDs","Conditional_NS_IDs","Non_Test_Generating_NS_IDs"]:
            if col in hdr: primary_ids.update(split_ids(primary.cell(r,hdr[col]).value))
    all_rtm_ids=rtmids|primary_ids
    ck("rtm_and_normative_statement_identifier_sets_equal",nsids==all_rtm_ids,{"rtm_unique_identifiers":len(all_rtm_ids),"document_unique_identifiers":len(nsids),"only_in_rtm":sorted(all_rtm_ids-nsids),"only_in_document":sorted(nsids-all_rtm_ids)})
    ds=wb["NS_Inventory_Dispositions"]; dh={ds.cell(1,c).value:c for c in range(1,ds.max_column+1)}; rec={k:ds.cell(2,c).value for k,c in dh.items()}
    disposition_ok=(rec.get("Unassigned_Alias")==ALIAS and rec.get("Canonical_NS_ID")==CANONICAL and rec.get("Disposition")=="CANONICALIZED_DUPLICATE" and rec.get("Unique_NS_Identifier_Count")==EXPECTED_UNIQUE and rec.get("Core_Normative_Occurrence_Count")==EXPECTED_OCCURRENCES and rec.get("Specification_Version")==RTM_SPEC_VERSION)
    ck("rtm_inventory_disposition_is_controlled",disposition_ok,rec)
    hist=(ROOT/"RELEASE_NOTES_v2.0.0.md").read_text(encoding="utf-8")
    ck("historical_v2_0_0_count_claim_preserved","358 Core-derived atomic statements" in hist,"RELEASE_NOTES_v2.0.0.md retains historical wording")
    notes=(ROOT/CURRENT_RELEASE_NOTES).read_text(encoding="utf-8")
    ck("current_release_notes_record_inventory_correction",all(x in notes for x in ["357 unique controlled Normative Statement identifiers","358 controlled normative source-text occurrences","NS-17.2-01","NS-2.7-01"]),CURRENT_RELEASE_NOTES)
    readme=(ROOT/"README.md").read_text(encoding="utf-8")
    ck("readme_records_current_inventory",all(x in readme for x in ["357 unique Normative Statement identifiers","358 controlled normative source-text occurrences","NS-17.2-01"]),"README.md")
    ck("controlled_count_reconciliation",len(nsids)+max(0,len(core_hits)-1)==EXPECTED_OCCURRENCES,{"unique_ns_identifier_count":len(nsids),"duplicate_core_occurrences_beyond_canonical":max(0,len(core_hits)-1),"core_normative_occurrence_count":len(nsids)+max(0,len(core_hits)-1)})
    sources=["spec/AGCP-Core.docx","spec/AGCP Normative Statements.docx","spec/AGCP_Requirements_Traceability_Matrix_(RTM).xlsx","RELEASE_NOTES_v2.0.0.md",CURRENT_RELEASE_NOTES,"README.md","governance/validate_normative_statement_inventory.py"]
    report={"release_context":release_context(),"validation_type":"AGCP_NORMATIVE_STATEMENT_INVENTORY_VALIDATION","status":"PASS" if not issues else "FAIL","inventory":{"unique_ns_identifier_count":EXPECTED_UNIQUE,"core_normative_occurrence_count":EXPECTED_OCCURRENCES,"canonical_ns_id":CANONICAL,"unassigned_alias":ALIAS,"duplicate_core_sections":["2.7","17.2"],"disposition":"CANONICALIZED_DUPLICATE","reason":"The Section 17.2 source-text occurrence duplicates the obligation already identified as NS-2.7-01 and does not create a second independently testable obligation."},"checks_passed":sum(c["status"]=="PASS" for c in checks),"checks_total":len(checks),"checks":checks,"issues":issues,"source_hashes":{r:sha(ROOT/r) for r in sources}}
    REPORT.write_text(json.dumps(report,indent=2,ensure_ascii=False)+"\n",encoding="utf-8")
    print(json.dumps({"status":report["status"],"checks_passed":report["checks_passed"],"checks_total":report["checks_total"],"issues":issues},indent=2))
    return 0 if not issues else 1
if __name__=="__main__": raise SystemExit(main())
